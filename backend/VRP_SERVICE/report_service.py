from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
from typing import Tuple
from datetime import datetime

from .export_paths import EXPORTS_DIR, LOGOS_DIR
from backend.VRP_DATABASE.database import get_conn

LOGO_PATH = LOGOS_DIR / "NOVAES.png"

# ---------- helpers de campos ----------
def _add_field(paragraph, field_code: str, visible_hint: str = ""):
    """Insere um campo Word (TOC/LOF/etc)."""
    r = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = field_code
    fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
    r._r.append(fld_begin); r._r.append(instr); r._r.append(fld_sep)
    if visible_hint:
        paragraph.add_run(visible_hint)
    fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(fld_end)

def _add_seq_figure_caption(paragraph, caption_text: str, label: str = "Figura"):
    """Legenda com numeração automática SEQ."""
    paragraph.style = "Caption"
    paragraph.add_run(f"{label} ")
    r = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = f"SEQ {label} \\* ARABIC"
    s = OxmlElement("w:fldChar"); s.set(qn("w:fldCharType"), "separate")
    r._r.append(b); r._r.append(instr); r._r.append(s)
    paragraph.add_run("1")
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(e)
    paragraph.add_run(f": {caption_text}")

# ---------- salvamento robusto ----------
def _safe_save_docx(doc: Document, target: Path) -> Path:
    """Salva DOCX; se bloqueado, salva com timestamp."""
    target.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(target)
        return target
    except PermissionError:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        alt = target.with_name(f"{target.stem}_{ts}.docx")
        doc.save(alt)
        return alt

def _next_pdf_path_for(docx_path: Path) -> Path:
    """Sugere caminho PDF; se bloqueado, usa timestamp."""
    base_pdf = docx_path.with_suffix(".pdf")
    try:
        if base_pdf.exists():
            base_pdf.unlink()
        return base_pdf
    except PermissionError:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        return docx_path.with_name(f"{docx_path.stem}_{ts}.pdf")

# ---------- formatação ----------
def _set_default_fonts(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

def _month_pt_br_upper(dt: datetime) -> str:
    meses = ["JANEIRO","FEVEREIRO","MARÇO","ABRIL","MAIO","JUNHO","JULHO","AGOSTO","SETEMBRO","OUTUBRO","NOVEMBRO","DEZEMBRO"]
    return f"{meses[dt.month-1]} – {dt.year}"

def _add_header_logo(doc: Document):
    """Logo no cabeçalho de todas as seções."""
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = p.add_run()
        try:
            run.add_picture(str(LOGO_PATH), width=Cm(5.5))
        except Exception:
            p.add_run("[Logo NOVAES.png não encontrado]")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def _style_table(table):
    """Aplica estilo de grade e cabeçalho em negrito."""
    table.style = "Table Grid"
    # primeira linha como cabeçalho se fizer sentido
    if table.rows:
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs or [cell.paragraphs[0].add_run(cell.paragraphs[0].text)]:
                run.bold = True

# ---------- dados ----------
def _fetch_all(checklist_id: int):
    conn = get_conn()
    ck_row = conn.execute("SELECT * FROM checklists WHERE id=?", (checklist_id,)).fetchone()
    site_row = None
    if ck_row and ck_row["vrp_site_id"]:
        site_row = conn.execute("SELECT * FROM vrp_sites WHERE id=?", (ck_row["vrp_site_id"],)).fetchone()
    photos_rows = conn.execute(
        "SELECT * FROM photos WHERE checklist_id=? AND include_in_report=1 ORDER BY display_order,id",
        (checklist_id,)
    ).fetchall()
    conn.close()
    ck = dict(ck_row) if ck_row else {}
    site = dict(site_row) if site_row else {}
    photos = [dict(p) for p in photos_rows]
    return ck, site, photos

# ---------- página 4: introdução + Tabela 1 ----------
def _add_intro_and_table(doc: Document):
    # Texto mais conciso
    intro = (
        "Esta planilha consolida 18 VRPs do município de Maceió, com identificação, DN e link georreferenciado para cada ponto. "
        "A relação padroniza o acompanhamento operacional e subsidia inspeções, manutenções e auditorias."
    )
    p = doc.add_paragraph(intro)
    p.paragraph_format.first_line_indent = Cm(1)

    doc.add_paragraph(" ")
    doc.add_paragraph("Tabela 1: Planilha geral das localizações e descrições técnicas de cada VRP")

    data = [
        ("DMC-José Sampaio Luz","https://maps.app.goo.gl/EkDi2C8yQKDjAVrCA","300","CLA-VAL"),
        ("DMC-Carlos Tenório","https://maps.app.goo.gl/qDCvgPVvCgbvEVqu9","200","CLA-VAL"),
        ("DMC-Pio XII","https://maps.app.goo.gl/GqqWwuUht66kmVAE7","200","CLA-VAL"),
        ("DMC-José Guilherme","https://maps.app.goo.gl/h3f6zH6oKT59Yicj8","250","CLA-VAL"),
        ("DMC-José Lajes","https://maps.app.goo.gl/R55GYMHyFc1VCSaEA","250","CLA-VAL"),
        ("DMC-Salvador Calmon","https://maps.app.goo.gl/9g45ghpF2VhKRWGf7","200","CLA-VAL"),
        ("DMC-José Carneiro","https://maps.app.goo.gl/wDa8efRQgsycJWtJ9","150","CLA-VAL"),
        ("DMC-Inácio Gracindo","https://maps.app.goo.gl/kpZcFRNWF33jBE2E9","200","CLA-VAL"),
        ("DMC-Paulina Mendonça","https://maps.app.goo.gl/iubsuSsSgDudYMWS6","250","CLA-VAL"),
        ("DMC-Gustavo Paiva","https://maps.app.goo.gl/cHn5Yw485ZRhK66L7","200","CLA-VAL"),
        ("DMC-Dona Constança","https://maps.app.goo.gl/TcCnho57MvwortP47","200","CLA-VAL"),
        ("DMC-Pretestato Ferreira","https://maps.app.goo.gl/dXPgFUwb8FqfK7N66","150","CLA-VAL"),
        ("DMC-Aloísio Branco","https://maps.app.goo.gl/twMMzviBRXXdTcZi7","150","CLA-VAL"),
        ("DMC-Praça da Bíblia","https://maps.app.goo.gl/crAzBpJuC3tbiPyU8","150","CLA-VAL"),
        ("DMC-Travessa Nazaré","https://maps.app.goo.gl/twMMzviBRXXdTcZi7","200","CLA-VAL"),
        ("DMC-Benedito_Bentes_I","https://maps.app.goo.gl/EECnzX2u342UT3AR9","300","CLA-VAL"),
        ("DMC-Benedito_Bentes_II","https://maps.app.goo.gl/nmURQA4vqRqB5Wn49","200","CLA-VAL"),
        ("DMC-Recantos","https://maps.app.goo.gl/YzwzdrUm7sYGThVJ6","200","CLA-VAL"),
    ]
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Descrição"; hdr[1].text = "Localização"; hdr[2].text = "DN"; hdr[3].text = "Fabricante"
    for desc, link, dn, fab in data:
        r = table.add_row().cells
        r[0].text = desc; r[1].text = link; r[2].text = dn; r[3].text = fab
    _style_table(table)

# ---------- DOCX ----------
def build_docx(checklist_id: int, ai_text: str, pasta_destino: str = None) -> Path:
    ck, site, photos = _fetch_all(checklist_id)
    if pasta_destino:
        export_folder = Path(pasta_destino) / f"{checklist_id}"
    else:
        export_folder = EXPORTS_DIR / f"{checklist_id}"
    export_folder.mkdir(parents=True, exist_ok=True)
    fname = export_folder / f"Relatorio_VRP_{checklist_id}.docx"

    doc = Document()
    _set_default_fonts(doc)
    _add_header_logo(doc)

    # CAPA
    doc.add_paragraph(" ")
    title = doc.add_paragraph("RELATÓRIO DE ATIVIDADES EXECUTADAS")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.runs[0]; tr.bold = True; tr.font.size = Pt(26); tr.font.name = "Times New Roman"

    doc.add_paragraph(" ")
    objeto = doc.add_paragraph("OBJETO: AFERIÇÃO DE MACROMEDIDORES, CONTROLE E MANUTENÇÃO DE VRPS NA REGIÃO METROPOLITANA DE MACEIÓ – RMM.")
    objeto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    oru = objeto.runs[0]; oru.font.size = Pt(15); oru.font.name = "Times New Roman"

    doc.add_paragraph(" ")
    report_number = f"R{int(ck.get('id', checklist_id)):02d}" if ck.get("id") else f"R{int(checklist_id):02d}"
    mid = doc.add_paragraph(f"RELATÓRIO DE ATIVIDADES – {report_number}")
    mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mid.runs[0].font.size = Pt(12); mid.runs[0].font.name = "Times New Roman"

    # mês/ano
    try:
        dt = datetime.strptime(ck.get("date", ""), "%Y-%m-%d")
    except Exception:
        dt = datetime.now()
    bottom = doc.add_paragraph(_month_pt_br_upper(dt))
    bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    b = bottom.runs[0]; b.bold = True; b.font.size = Pt(12); b.font.name = "Times New Roman"

    doc.add_page_break()

    # p.2 SUMÁRIO
    doc.add_heading("SUMÁRIO", level=1)
    _add_field(doc.add_paragraph(), r'TOC \o "1-3" \h \z \u', " (Atualize com F9)")
    doc.add_page_break()

    # p.3 LISTA DE FIGURAS
    doc.add_heading("LISTA DE FIGURAS", level=1)
    _add_field(doc.add_paragraph(), r'TOC \h \z \c "Figura"', " (Atualize com F9)")
    _add_field(doc.add_paragraph(), r'TOC \h \z \c "Figure"')
    doc.add_page_break()

    # p.4 INTRODUÇÃO + Tabela 1
    doc.add_heading("INTRODUÇÃO", level=1)
    _add_intro_and_table(doc)
    doc.add_page_break()

    # p.5+ CONTEÚDO
    # Dados Técnicos (TABELA)
    doc.add_heading("Dados Técnicos da VRP", level=1)
    t1 = doc.add_table(rows=6, cols=2)
    drows = [
        ("Cidade", site.get("city", "")),
        ("Local", site.get("place", "")),
        ("Marca", site.get("brand", "")),
        ("Tipo", site.get("type", "")),
        ("DN (mm)", str(site.get("dn", ""))),
        ("Acesso / Tráfego / Tampas", f"{site.get('access_install','')} / {site.get('traffic','')} / {site.get('lids','')}")
    ]
    for i, (k, v) in enumerate(drows):
        t1.cell(i, 0).text = k
        t1.cell(i, 1).text = str(v)
    _style_table(t1)

    # Análise Hidráulica (TABELA) – sem observações do usuário
    doc.add_heading("Análise Hidráulica", level=1)
    t2 = doc.add_table(rows=3, cols=2)
    t2.cell(0, 0).text, t2.cell(0, 1).text = "Montante c/ registro", "Sim" if ck.get("has_reg_upstream") else "Não"
    t2.cell(1, 0).text, t2.cell(1, 1).text = "Jusante c/ registro", "Sim" if ck.get("has_reg_downstream") else "Não"
    t2.cell(2, 0).text, t2.cell(2, 1).text = "Bypass", "Sim" if ck.get("has_bypass") else "Não"
    _style_table(t2)

    # Análise de Pressão (mca) – TABELA
    doc.add_heading("Análise de Pressão (mca)", level=1)
    def n(v):
        try: return f"{float(v):.1f}"
        except: return "-"
    t3 = doc.add_table(rows=3, cols=3)
    t3.cell(0,0).text, t3.cell(0,1).text, t3.cell(0,2).text = "", "Antes (mca)", "Depois (mca)"
    t3.cell(1,0).text, t3.cell(1,1).text, t3.cell(1,2).text = "Montante", n(ck.get("p_up_before")), n(ck.get("p_up_after"))
    t3.cell(2,0).text, t3.cell(2,1).text, t3.cell(2,2).text = "Jusante",  n(ck.get("p_down_before")), n(ck.get("p_down_after"))
    _style_table(t3)

    # Análise Técnica (IA) – concisa (texto já vem sintetizado pela IA)
    doc.add_heading("Análise Técnica (IA)", level=1)
    doc.add_paragraph(ai_text or "—")

    # Figuras – 10 cm (alt) x 7,5 cm (larg) – legenda apenas com rótulo
    if photos:
        doc.add_heading("Figuras", level=1)
        for ph in photos:
            path = ph.get("file_path")
            if not path:
                continue
            try:
                doc.add_picture(path, width=Cm(7.5), height=Cm(10))
            except Exception:
                doc.add_paragraph(f"[Falha ao inserir: {path}]")
            cap = doc.add_paragraph("")
            _add_seq_figure_caption(cap, f"{ph.get('label','')}", label="Figura")

    # salva de forma resiliente
    fname = _safe_save_docx(doc, fname)
    return fname

def convert_to_pdf(docx_path: Path) -> Path | None:
    """DOCX -> PDF (usa nome alternativo se arquivo estiver bloqueado)."""
    try:
        from docx2pdf import convert
        out = _next_pdf_path_for(docx_path)
        convert(str(docx_path), str(out))
        return out
    except Exception:
        return None

def generate_full_report(checklist_id: int, ai_text: str, pasta_destino: str = None) -> Tuple[str, str | None]:
    docx_path = build_docx(checklist_id, ai_text, pasta_destino)
    pdf_path = convert_to_pdf(docx_path)

    conn = get_conn()
    conn.execute("""
        INSERT INTO reports (checklist_id, ai_summary, docx_path, pdf_path)
        VALUES (?,?,?,?)
        ON CONFLICT(checklist_id) DO UPDATE SET
            ai_summary=excluded.ai_summary,
            docx_path=excluded.docx_path,
            pdf_path=excluded.pdf_path
    """, (checklist_id, ai_text, str(docx_path), str(pdf_path) if pdf_path else None))
    conn.commit(); conn.close()
    return str(docx_path), (str(pdf_path) if pdf_path else None)
